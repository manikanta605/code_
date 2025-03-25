from fastapi import FastAPI, HTTPException ,UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pdf2image.pdf2image import convert_from_path
from pydantic import BaseModel
import os
import shutil

from datetime import datetime

from db import get_db_connection
from exceltotext import exceltotext
from imagetoText import imageToText, insert_excel_row
from pdf2image_convert import insert_image_classification, insert_pdf_conversion, process_pdfs

import pandas as pd
from io import BytesIO
import openpyxl  # Add this import
import psycopg2  # Add this import
import docx  # Add this import
import extract_msg  # Add this import

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # You can restrict this to specific origins like ["http://localhost:4200"]
    allow_credentials=True,
    allow_methods=["*"],  # Allow specific methods like ["GET", "POST"]
    allow_headers=["*"],  # Allow specific headers
)


class FolderPath(BaseModel):
    path: str


class FolderPathRequest(BaseModel):
    root_input: str

class labelCreationdetials(BaseModel):
    label:str
    color:str
    
class jsonCreationdetials(BaseModel):
    file_name:str
    json_text_1:str
    json_text_2:str
    folder_name:str
    file_id:str



@app.post("/add_label")
async def process_folder(request:labelCreationdetials):
    try:
        insert_label_create(request.label, request.color)
        return {
            "message": "Lable created successfully.",
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")


@app.post("/add_model_json")
async def add_model_json(request:jsonCreationdetials):
    try:
        insert_model_json_create(request.file_name, request.json_text_1, request.json_text_2, request.folder_name,request.file_id)
        return {
            "message": "Lable created successfully.",
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")



@app.get("/get_label")
async def process_folder():
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        query = "SELECT * FROM tbl_label"
        cursor.execute(query)
        records = cursor.fetchall()
        colnames = [desc[0] for desc in cursor.description]
        result = [dict(zip(colnames, record)) for record in records]
        return result
        
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")


@app.get("/get_str_excel")
async def process_folder():
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        query = "SELECT sheet_name, text FROM str_data_excel"
        cursor.execute(query)
        records = cursor.fetchall()
        
        colnames = [desc[0] for desc in cursor.description]
        result = [dict(zip(colnames, record)) for record in records]
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")
    finally:
        cursor.close()
        conn.close()
        
@app.get("/image_data_text")
async def process_folder():
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        query = "SELECT * FROM image_data_text"
        cursor.execute(query)
        records = cursor.fetchall()
        
        colnames = [desc[0] for desc in cursor.description]
        result = [dict(zip(colnames, record)) for record in records]
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")
    finally:
        cursor.close()
        conn.close()

@app.get("/image_data_text/{id}")
async def process_folder(id:str):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        query = "SELECT * FROM image_data_text  where id = '"+id+"';"
        cursor.execute(query)
        records = cursor.fetchall()
        
        colnames = [desc[0] for desc in cursor.description]
        result = [dict(zip(colnames, record)) for record in records]
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")
    finally:
        cursor.close()
        conn.close()


# TRAIN_DATA_RAW = [
#     ("John Doe has Type 2 Diabetes and Hypertension.", ["Type 2 Diabetes", "Hypertension"], ["DIAGNOSIS", "DIAGNOSIS"]),
#     ("The patient suffers from severe chest pain and shortness of breath.", ["chest pain", "shortness of breath"], ["SYMPTOM", "SYMPTOM"]),
#     ("He is taking Metformin, Amlodipine, and Atorvastatin.", ["Metformin", "Amlodipine", "Atorvastatin"], ["MEDICATION", "MEDICATION", "MEDICATION"]),
#     ("Lab tests show high LDL Cholesterol and elevated Troponin I.", ["LDL Cholesterol", "Troponin I"], ["TEST", "TEST"]),
#     ("Diagnosis: Unstable Angina, Hypertension, Type 2 Diabetes.", ["Unstable Angina", "Hypertension", "Type 2 Diabetes"], ["DIAGNOSIS", "DIAGNOSIS", "DIAGNOSIS"]),
# ]



# Function to insert PDF conversion results into the database
def insert_label_create(label, color):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute(
            "INSERT INTO tbl_label ( label, color) VALUES (%s, %s)",
            (label, color)
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"Error inserting PDF conversion data: {str(e)}")
    finally:
        cursor.close()
        conn.close()
        
        

@app.post("/pdf2image")
async def process_folder(request: FolderPathRequest):
    
    try:
        root_input = request.root_input
        root_output = os.path.join(os.getcwd(), "output_image")
        
        if os.path.exists(root_output):
            shutil.rmtree(root_output)
        os.makedirs(root_output, exist_ok=True)
        # return
        print("step 1 started ------------>")
        start_time = datetime.now()
        # Step 1: Convert PDFs to images
        pdf_conversion_result = process_pdfs(root_input, root_output)
        if not pdf_conversion_result:
            raise HTTPException(status_code=400, detail="PDF conversion failed or no PDFs found.")
        end_time = datetime.now()
        print('Duration: {}'.format(end_time - start_time))
        return {
            "message": "Process completed successfully.",
            # "pdf_conversion": pdf_conversion_result,
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")
    



@app.post("/imagetoText")
async def process_folder():
    try:
        imageToText()
        return {
            "message": "Process completed successfully.",
        }
        
    except TypeError as e:
        raise HTTPException(status_code=500, detail=f"Type error occurred during the process: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")



@app.post("/upload-excel/")
async def upload_excel(request: FolderPathRequest):
    dirpath_full = []
    try:
        root_input = request.root_input
        skipped_files = []
        output_folder = os.path.join(os.getcwd(), "output_excel_txt")
        
        if os.path.exists(output_folder):
            shutil.rmtree(output_folder)
        os.makedirs(output_folder, exist_ok=True)
        
        for dirpath, dirnames, filenames in os.walk(root_input):
            dirpath_full.append(dirnames)
            files = [f for f in filenames if f.endswith('.xlsx') or f.endswith('.xls')]
            print(files, 'files')
            
            for file_name in files:
                try:
                    print(file_name, 'file_name')
                    file_path = os.path.join(dirpath, file_name)
                    with open(file_path, 'rb') as f:
                        contents = f.read()
                    excel_data = pd.ExcelFile(BytesIO(contents))

                    for sheet_name in excel_data.sheet_names:
                        df = pd.read_excel(excel_data, sheet_name=sheet_name)
                        
                        # Add headers as a separate record
                        headers = ','.join(df.columns)
                        print("headers", sheet_name, headers)
                        
                        sheet_output_path = os.path.join(output_folder, f"{file_name}_{sheet_name}.txt")
                        with open(sheet_output_path, 'w') as txt_file:
                            txt_file.write(headers + '\n')
                            for _, row in df.iterrows():
                                concatenated_columns = ','.join(map(str, row.values))
                                txt_file.write(concatenated_columns + '\n')
                                
                                result_1 = "/".join(filter(None, [item[0] if item else '' for item in dirpath_full]))
                                print("result_1", result_1)
                                print("concatenated_columns", sheet_name, concatenated_columns)
                                insert_excel_row(file_name, sheet_name, headers, concatenated_columns, result_1)
                except Exception as e:
                    skipped_files.append(file_path)
                    print(f"Skipped {file_path} due to error: {str(e)}")

        # Log skipped files
        if skipped_files:
            log_file_path = os.path.join(root_input, "skipped_files.log")
            with open(log_file_path, 'w') as log_file:
                for skipped_file in skipped_files:
                    log_file.write(f"{skipped_file}\n")
            print(f"Skipped files logged in {log_file_path}")

        return {"message": "Excel files processed successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")

@app.get("/json_data_text/{id}")
async def process_folder(id:str):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        query = "SELECT * FROM json_data_str  where file_id = '"+id+"';"
        cursor.execute(query)
        records = cursor.fetchall()
        
        colnames = [desc[0] for desc in cursor.description]
        result = [dict(zip(colnames, record)) for record in records]
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")
    finally:
        cursor.close()
        conn.close()


# Function to insert PDF conversion results into the database
def insert_model_json_create(file_name,json_text_1,json_text_2, folder_name, file_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute(
            "INSERT INTO json_data_str ( file_name, json_text_1, json_text_2, folder_name, file_id) VALUES (%s, %s,%s, %s, %s)",
            (file_name,json_text_1,json_text_2, folder_name,file_id)
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"Error inserting JSON data: {str(e)}")
    finally:
        cursor.close()
        conn.close()

@app.post("/upload-doc/")
async def upload_doc(file: UploadFile = File(...)):
    try:
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="Invalid file format. Only .docx files are supported.")
        
        contents = await file.read()
        doc = docx.Document(BytesIO(contents))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        document_text = '\n'.join(full_text)
        
        # Insert document text into the database
        insert_document_text(file.filename, document_text)
        
        return {"message": "Document file processed and inserted into the database successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")

# Function to insert document text into the database
def insert_document_text(file_name, document_text):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute(
            "INSERT INTO document_data (file_name, document_text) VALUES (%s, %s)",
            (file_name, document_text)
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"Error inserting document data: {str(e)}")
    finally:
        cursor.close()
        conn.close()
        
@app.post('/test_all_files')
async def test_all_files(request: FolderPathRequest):
    dirpath_full = []
    try:
        root_input = request.root_input
        skipped_files = []
        output_folder = os.path.join(os.getcwd(), "output_files")
        
        if os.path.exists(output_folder):
            shutil.rmtree(output_folder)
        os.makedirs(output_folder, exist_ok=True)
        
        for dirpath, dirnames, filenames in os.walk(root_input):
            dirpath_full.append(dirnames)
            for file_name in filenames:
                file_path = os.path.join(dirpath, file_name)
                try:
                    if file_name.endswith('.xlsx') or file_name.endswith('.xls'):
                        # Process Excel files
                        with open(file_path, 'rb') as f:
                            contents = f.read()
                        excel_data = pd.ExcelFile(BytesIO(contents))

                        for sheet_name in excel_data.sheet_names:
                            df = pd.read_excel(excel_data, sheet_name=sheet_name)
                            
                            # Add headers as a separate record
                            headers = ','.join(df.columns)
                            print("headers", sheet_name, headers)
                            
                            sheet_output_path = os.path.join(output_folder, f"{file_name}_{sheet_name}.txt")
                            with open(sheet_output_path, 'w') as txt_file:
                                txt_file.write(headers + '\n')
                                for _, row in df.iterrows():
                                    concatenated_columns = ','.join(map(str, row.values))
                                    txt_file.write(concatenated_columns + '\n')
                                    
                                    result_1 = "/".join(filter(None, [item[0] if item else '' for item in dirpath_full]))
                                    print("result_1", result_1)
                                    print("concatenated_columns", sheet_name, concatenated_columns)
                                    insert_excel_row(file_name, sheet_name, headers, concatenated_columns, result_1)
                    elif file_name.endswith('.docx'):
                        # Process DOCX files
                        with open(file_path, 'rb') as f:
                            contents = f.read()
                        doc = docx.Document(BytesIO(contents))
                        full_text = []
                        for para in doc.paragraphs:
                            full_text.append(para.text)
                        document_text = '\n'.join(full_text)
                        print("document_text", document_text)
                        
                        # Create a .txt file for the .docx content
                        doc_output_path = os.path.join(output_folder, f"{file_name}.txt")
                        with open(doc_output_path, 'w') as txt_file:
                            txt_file.write(document_text)
                        
                        # Insert document text into the database
                        insert_document_text(file_name, document_text)
                    elif file_name.endswith('.pdf'):
                        # Process PDF files
                        pdf_path = file_path
                        pdf_filename = file_name
                        print(pdf_filename, 'pdf_filename')
                        try:
                            images = convert_from_path(pdf_path)
                            for i, image in enumerate(images, start=1):
                                image_filename = f"{pdf_filename}_{i}.png"
                                image_path = os.path.join(output_folder, image_filename)
                                result_1 = "/".join(filter(None, [item[0] if item else '' for item in dirpath_full]))

                                print("result_1", result_1)
                                insert_image_classification(pdf_filename, image_filename, image_path, result_1)
                                image.save(image_path, "PNG")
                        except Exception as e:
                            print(f"Error processing {pdf_filename}: {e}")
                            continue  # Skip to the next PDF file

                        insert_pdf_conversion(result_1, result_1, pdf_filename, '.pdf', len(images), "Success")
                    elif file_name.endswith('.msg'):
                        # Process MSG files
                        msg = extract_msg.Message(file_path)
                        msg_text = msg.body
                        print("msg_text", msg_text)
                        
                        # Create a .txt file for the .msg content
                        msg_output_path = os.path.join(output_folder, f"{file_name}.txt")
                        with open(msg_output_path, 'w') as txt_file:
                            txt_file.write(msg_text)
                        
                        # Insert MSG text into the database
                        insert_msg_text(file_name, msg_text)
                except Exception as e:
                    skipped_files.append(file_path)
                    print(f"Skipped {file_path} due to error: {str(e)}")

        # Log skipped files
        if skipped_files:
            log_file_path = os.path.join(root_input, "skipped_files.log")
            with open(log_file_path, 'w') as log_file:
                for skipped_file in skipped_files:
                    log_file.write(f"{skipped_file}\n")
            print(f"Skipped files logged in {log_file_path}")

        return {"message": "Files processed successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during the process: {str(e)}")

# Function to insert MSG text into the database
def insert_msg_text(file_name, msg_text):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute(
            "INSERT INTO msg_data (file_name, msg_text) VALUES (%s, %s)",
            (file_name, msg_text)
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"Error inserting MSG data: {str(e)}")
    finally:
        cursor.close()
        conn.close()
